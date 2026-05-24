"""
Contract filling engine.
Takes form data and fills it into the vehicle pledge loan contract template.
Outputs a ready-to-print .docx file.
"""
from docx import Document
from pathlib import Path
import re, os, copy

TEMPLATE_PATH = Path(__file__).parent / "static" / "contract" / "template.docx"


def fill_contract(data: dict) -> bytes:
    """
    data keys:
      borrower_name, borrower_id, borrower_phone, borrower_addr
      lender_name, lender_id, lender_phone, lender_addr
      loan_amount_cn, loan_amount_num, loan_term, loan_date_start, loan_date_end
      monthly_rate, interest_pay_day, loan_purpose
      vehicle_brand, vehicle_plate, vehicle_engine, vehicle_vin, vehicle_reg, vehicle_color
      bank_name, bank_branch, bank_account
      iou_date
    """
    doc = Document(str(TEMPLATE_PATH))

    replacements = {
        "借款人：                       ": f"借款人：{data.get('borrower_name', '')}" + " " * 30,
        "身份证号码：                           ": f"身份证号码：{data.get('borrower_id', '')}",
        "电话：                         ": f"电话：{data.get('borrower_phone', '')}",
        "地址：                                 。": f"地址：{data.get('borrower_addr', '')}。",
        "出借人：                       ": f"出借人：{data.get('lender_name', '')}",
        "身份证号码：                           ": f"身份证号码：{data.get('lender_id', '')}",
        "电话：                         ": f"电话：{data.get('lender_phone', '')}",
        "地址：                                 。": f"地址：{data.get('lender_addr', '')}。",
        "借款金额为人民币（大写）            元（小写¥       元）": (
            f"借款金额为人民币（大写）{data.get('loan_amount_cn', '')}元（小写¥{data.get('loan_amount_num', '')}元）"
        ),
        "现金借款（大写）          元人民币（小写：¥       元）": (
            f"现金借款（大写）{data.get('loan_amount_cn', '')}元人民币（小写：¥{data.get('loan_amount_num', '')}元）"
        ),
        "    个月，自    年   月": f"    {data.get('loan_term', '')}个月，自{data.get('loan_date_start', '')}",
        "日至    年   月   日止": f"日至{data.get('loan_date_end', '')}止",
        "自   年   月   日至    年   月   日止": (
            f"自{data.get('loan_date_start', '')}至{data.get('loan_date_end', '')}止"
        ),
        "借款期限    个月": f"借款期限{data.get('loan_term', '')}个月",
        "月利率为     %": f"月利率为{data.get('monthly_rate', '')}%",
        "利息按月利率   %计算": f"利息按月利率{data.get('monthly_rate', '')}%计算",
        "利息支付日为每月    日": f"利息支付日为每月{data.get('interest_pay_day', '')}日",
        "借款人的借款用途为          ": f"借款人的借款用途为{data.get('loan_purpose', '')}",
        "品牌型号：                         ": f"品牌型号：{data.get('vehicle_brand', '')}",
        "号牌号码：                         ": f"号牌号码：{data.get('vehicle_plate', '')}",
        "发动机号码：                       ": f"发动机号码：{data.get('vehicle_engine', '')}",
        "车架号：                          ": f"车架号：{data.get('vehicle_vin', '')}",
        "登记证号：                         ": f"登记证号：{data.get('vehicle_reg', '')}",
        "车身颜色：                         ": f"车身颜色：{data.get('vehicle_color', '')}",
        "户名：            ": f"户名：{data.get('bank_name', '')}",
        "开户行：                 ": f"开户行：{data.get('bank_branch', '')}",
        "卡号：                            ": f"卡号：{data.get('bank_account', '')}",
        "户名：            开户行：                   卡号": (
            f"户名：{data.get('bank_name', '')} 开户行：{data.get('bank_branch', '')} 卡号：{data.get('bank_account', '')}"
        ),
        "卡号                                ": f"卡号{data.get('bank_account', '')}",
        "    年   月   日": f"  {data.get('iou_date', '')}",
        "签订日期：        年     月     日": f"签订日期：{data.get('iou_date', '')}",
        "日期：     年   月    日": f"日期：{data.get('iou_date', '')}",
    }

    # Replace in paragraphs
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                replace_in_paragraph(para, old, new)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            replace_in_paragraph(para, old, new)

    # Save to bytes
    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving formatting."""
    full_text = paragraph.text
    if old_text not in full_text:
        return

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # Set the first run to the new text
    if paragraph.runs:
        paragraph.runs[0].text = full_text.replace(old_text, new_text)
    else:
        # No runs - add a new one
        from docx.oxml.ns import qn
        new_run = paragraph.add_run(full_text.replace(old_text, new_text))
        # Copy style from original if available
        if paragraph.style:
            new_run.font.size = paragraph.style.font.size
